import re
import docx

# Helper function to check if two strings are transposition errors
def is_transposition(word1, word2):
    if len(word1) != len(word2):
        return False
    diffs = [(c1, c2) for c1, c2 in zip(word1, word2) if c1 != c2]
    return len(diffs) == 2 and diffs[0] == diffs[1][::-1]

# Function to apply bold and color to a word in a paragraph
def apply_formatting(paragraph, word, color, bold=True):
    run = paragraph.add_run(word + " ")
    run.bold = bold
    run.font.color.rgb = docx.shared.RGBColor(*color)

# Function to replace and format words
def replace_and_format_words(paragraph, input_dict):
    original_text = paragraph.text  # Get the original paragraph text
    
    # Split paragraph into words
    words = re.findall(r'\w+|\s+|[^\w\s]+', original_text)  # Keeps words, spaces, and punctuation separate

    formatted_text = []
    
    for w in words:
        replaced = False
        if w.strip():  # Avoid empty strings
            for word, details in input_dict.items():
                synonyms = [word] + details['synonyms']
                replacement_word = details['replacement']

                # Exact match with case-insensitive comparison
                if w.lower() in [syn.lower() for syn in synonyms]:
                    formatted_text.append((replacement_word, (0, 0, 139), True))  # Bold, dark blue
                    replaced = True
                    break

                # Check for typo (transposition)
                for synonym in synonyms:
                    if is_transposition(synonym.lower(), w.lower()):
                        formatted_text.append((replacement_word + "_typo", (255, 0, 0), True))  # Bold, red
                        replaced = True
                        break

                if replaced:
                    break

        if not replaced:
            formatted_text.append((w, (0, 0, 0), False))  # No replacement, regular black text

    # Clear the paragraph and rebuild it with formatted runs
    paragraph.clear()

    for word, color, bold in formatted_text:
        apply_formatting(paragraph, word, color, bold)

# Example usage
input_dict = {
    'happy': {
        'synonyms': ['joyful', 'content', 'glad'],
        'replacement': 'satisfied'
    },
    'quick': {
        'synonyms': ['fast', 'speedy', 'rapid'],
        'replacement': 'swift'
    }
}

doc = docx.Document('test.docx')

for paragraph in doc.paragraphs:
    replace_and_format_words(paragraph, input_dict)

doc.save('output.docx')
